from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = Path(r"F:/Mini-Drop/work-template.docx")
OUT = Path(
    r"C:/Users/Mozero/Documents/Codex/2026-06-12/https-www-anthropic-com-engineering-building/outputs/项目5-姓名-腾讯mini项目前置学习笔记.docx"
)


def set_cell_text(cell, text, font_size=10.5):
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)
    set_cell_vcenter(cell)


def set_cell_vcenter(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")


def main():
    doc = Document(TEMPLATE)
    table = doc.tables[0]

    # Basic info
    set_cell_text(table.rows[0].cells[1], "请填写")
    set_cell_text(table.rows[0].cells[5], "请填写")
    set_cell_text(table.rows[0].cells[8], "请填写")

    set_cell_text(table.rows[1].cells[2], "项目5：Mini-Drop 性能诊断系统复刻")

    choose_reason = (
        "我选择 Mini-Drop 项目，是因为它不是单一功能开发，而是一个完整的端到端工程系统。"
        "它包含 Web UI、Server、Agent、Analyzer、存储、状态机、心跳和可视化等多个模块，"
        "可以帮助我理解一个真实生产级性能诊断平台是如何拆分和协作的。\n"
        "同时，这个项目和 Linux 性能分析、自动化采集、火焰图、eBPF、智能归因等方向有关，"
        "能够锻炼我从需求理解、系统设计到工程落地的综合能力。"
    )
    set_cell_text(table.rows[2].cells[2], choose_reason)

    project_analysis = (
        "我认为这个项目的难点主要有四个。\n"
        "第一是端到端链路长。用户从 Web 创建任务后，需要经过 Server 调度、Agent 采集、"
        "结果上传、Analyzer 分析、Web 展示等多个环节，只要其中一个环节不稳定，整体演示就会失败。\n"
        "第二是 Agent 采集环境复杂。Agent 需要在 Linux 机器上采集目标进程性能数据，可能遇到 PID 不存在、"
        "权限不足、perf 不可用、采集超时等问题，因此必须有明确的错误处理和状态上报。\n"
        "第三是任务状态机和心跳机制要求清晰。题目要求每次状态迁移都要落库并带 reason 字段，"
        "Agent 还要定期心跳，Server 要能判断离线和恢复，这体现了生产系统的可观测性和可审计性。\n"
        "第四是分析结果需要可视化。原始性能数据不能直接给用户看，需要 Analyzer 转换成火焰图、"
        "热点函数列表或其他图表，这要求我理解采集数据到可视化结果之间的转换过程。"
    )
    set_cell_text(table.rows[3].cells[2], project_analysis)

    role_analysis = (
        "我在该项目中可以重点负责 Agent 侧和文档驱动开发相关工作。具体包括：理解 Agent 在系统中的职责，"
        "梳理 Agent 心跳、任务接收、采集执行、结果上报的流程；参与设计任务状态机和失败原因分类；"
        "整理题目要求、复刻指南和项目分析文档；在开发过程中把需求拆成可执行的 backlog；"
        "后续可以继续学习 perf 采集、火焰图生成和 eBPF 最小采集器的实现方式。\n"
        "同时，我也可以协助完成端到端演示链路的验证，例如检查从创建任务到生成火焰图的每一步状态是否正确，"
        "整理演示脚本和 README，保证项目在提交时可复现、可解释。"
    )
    set_cell_text(table.rows[4].cells[2], role_analysis)

    resources = (
        "1. Mini-Drop 项目题目文档：了解项目背景、基础能力、扩展能力、交付物和演示要求。\n"
        "2. Drop 系统复刻指南：学习 Web UI、API Server、Agent、Analyzer 四个模块的职责划分和端到端流程。\n"
        "3. Brendan Gregg Flame Graphs 文档：学习火焰图的用途和基本理解方式。\n"
        "4. Linux perf 相关资料：学习 perf record、采样频率、目标 PID、调用栈采样等基本概念。\n"
        "5. eBPF / bpftrace 入门资料：了解 eBPF 可用于观察 IO、调度等内核态事件。\n"
        "6. Anthropic《Building effective agents》：学习 workflow 与 agent 的区别，以及工具、流程和评估的重要性。\n"
        "7. OpenAI《A practical guide to building AI agents》：学习 Agent 的模型、工具、指令、编排和护栏设计思路。"
    )
    set_cell_text(table.rows[5].cells[2], resources)

    plan = (
        "6 月 14 日：已完成题目阅读和资料整理，初步理解 Mini-Drop 是一个 Web UI、Server、Agent、Analyzer 协作的性能诊断平台，并开始用文档驱动方式拆解项目。\n"
        "6 月 15 日：学习 Agent 的职责，包括心跳、任务接收、采集执行、结果上报；重点理解为什么 Agent 是系统中真正接触 Linux 采集环境的模块。\n"
        "6 月 16 日：学习 perf 基础，包括 PID、采样时长、采样率、调用栈采样和火焰图生成流程。\n"
        "6 月 17 日：学习任务状态机和 Agent 心跳机制，理解 PENDING -> RUNNING -> UPLOADING -> DONE / FAILED 的状态迁移，以及每次迁移记录 reason 的意义。\n"
        "6 月 18 日：学习 Analyzer 的职责，理解如何把原始采集数据转换成火焰图 SVG、TopN 热点函数和分析建议。\n"
        "6 月 19 日：学习 eBPF 的基本作用，重点了解它为什么适合观察 IO 或调度异常，以及在 Mini-Drop 中如何作为扩展采集器。\n"
        "6 月 20 日：整理学习笔记，补充学习心得，检查是否覆盖项目原因、项目难点、岗位职能、学习资料和学习计划。\n"
        "6 月 21 日：最终修改和提交。"
    )
    set_cell_text(table.rows[6].cells[2], plan)

    q1_problem = (
        "我一开始看 Mini-Drop 题目时，觉得它同时出现 Web UI、Server、Agent、Analyzer、perf、eBPF、"
        "火焰图、智能归因等很多概念，不清楚这个项目到底要做什么，也不知道应该先学哪一部分。"
    )
    set_cell_text(table.rows[8].cells[6], q1_problem)

    q1_understand = (
        "空：题目包含多个模块和多个技术关键词，信息密度很高。\n"
        "雨：如果直接从某个技术点开始学，例如一上来就学 eBPF，容易忽略项目主链路，导致学了很多细节但仍然不知道系统整体如何运行。\n"
        "伞：我认为应该先从整体流程入手，把项目理解成“Web 下发任务、Agent 采集数据、Analyzer 生成图表、Web 展示结果”的闭环，再逐步学习每个模块的职责。"
    )
    set_cell_text(table.rows[9].cells[6], q1_understand)

    q1_action = (
        "我先阅读了 Mini-Drop 题目和 Drop 系统复刻指南，并把项目拆成 Web UI、Server、Agent、Analyzer、Storage 几个模块。"
        "同时，我在项目目录中建立了文档驱动开发结构，整理了项目简介、MVP 范围、架构图、状态机和开发计划。"
    )
    set_cell_text(table.rows[10].cells[6], q1_action)

    q1_ai = (
        "使用工具：ChatGPT / Codex\n"
        "我的提问：Mini-Drop 这个题目我听不懂，到底要做什么？\n"
        "AI 的回答：它要做的是一个可以通过网页控制 Linux 性能采集，并把采集结果做成图表展示出来的小平台。"
        "用户在网页输入 PID、采样时长和采样率，Server 保存并下发任务，Agent 在 Linux 机器上执行采集，"
        "Analyzer 把原始数据变成火焰图和热点分析，最后 Web 展示结果。最核心的一句话是：网页下发任务，"
        "Agent 采集进程性能，Analyzer 生成图，网页展示结果。"
    )
    set_cell_text(table.rows[11].cells[6], q1_ai)

    q1_mentor = (
        "AI 的回答帮助我理解了项目主链路，但我还想向导师进一步确认：在提交前，基础能力和扩展能力之间应该如何取舍？"
        "如果时间有限，是优先保证端到端主链路稳定，还是优先多做几个扩展采集器？"
    )
    set_cell_text(table.rows[12].cells[6], q1_mentor)

    q2_problem = (
        "我不理解为什么项目中要单独设计 Agent，而不是让 Server 直接执行 perf 或写一个普通脚本来采集数据。"
    )
    set_cell_text(table.rows[13].cells[6], q2_problem)

    q2_understand = (
        "空：性能数据采集必须发生在目标 Linux 主机上，并且需要接触目标进程、系统权限、采集工具和本机环境。\n"
        "雨：如果 Server 直接采集，就要求所有目标进程都在 Server 所在机器上，不符合远程采集场景；"
        "如果只是普通脚本，也缺少心跳、任务状态、失败上报、超时控制和审计能力。\n"
        "伞：因此需要 Agent 部署在目标机器上，由 Agent 负责本地采集、心跳、任务执行和结果上报，Server 负责统一调度和管理。"
    )
    set_cell_text(table.rows[14].cells[6], q2_understand)

    q2_action = (
        "我对照题目要求，重点阅读了 Agent 心跳、Server 离线判定、任务状态机、采集器执行等内容，"
        "并把 Agent 的职责整理为：心跳、接收任务、执行采集、上传结果、上报状态。"
    )
    set_cell_text(table.rows[15].cells[6], q2_action)

    q2_ai = (
        "使用工具：ChatGPT / Codex\n"
        "我的提问：Mini-Drop 为什么需要 Agent，而不是普通 workflow 或脚本？\n"
        "AI 的回答：Agent 的价值在于它运行在目标机器上，能处理本地环境中的动态情况，例如 PID 是否存在、"
        "采集工具是否可用、权限是否足够、采集是否超时。普通脚本只能完成一次固定动作，而 Agent 可以持续心跳、"
        "接收任务、报告状态、处理失败，并和 Server 形成可观测的任务闭环。"
    )
    set_cell_text(table.rows[16].cells[6], q2_ai)

    q2_mentor = (
        "我认为 AI 的回答解决了我对 Agent 职责的基本疑问。但我还想进一步向导师确认：在 Mini-Drop 的简化实现中，"
        "Agent 是否必须做成长驻进程，还是可以先用定时轮询或命令行 worker 方式模拟 Agent，再逐步演进成长驻模式？"
    )
    set_cell_text(table.rows[17].cells[6], q2_mentor)

    # Apply a compact font size to long answer cells to reduce table overflow.
    for row_idx in range(2, 18):
        for paragraph in table.rows[row_idx].cells[-1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
